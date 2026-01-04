import pandas as pd
import matplotlib.pyplot as plt
import io
import xlsxwriter
import google.generativeai as genai
from docx import Document
from docx.shared import Inches
import streamlit as st # Need st for st.error, st.info, st.cache_data if used

# Use st.cache_data for performance if we were in handling streamlit, 
# but for pure utils we might want to avoid direct st calls if possible 
# to keep it clean. However, I used st.error in load_data. 
# I will modify load_data to return errors or print them, or pass a logger.
# For simplicity, I'll keep st dependency but check if it's running in streamlit?
# actually, st.error works even if not in streamlit? No, it prints to stderr or no-op?
# Let's keep st calls for now as it makes the app code simpler.

def load_data(file):
    """
    Loads CSV data, identifies relevant columns, and performs unit conversion.
    """
    # Determine file type and read
    try:
        if file.name.endswith('.xlsx') or file.name.endswith('.xls'):
            # It's an Excel file
            xlsx = pd.ExcelFile(file)
            
            # Smart Sheet Selection
            # 1. Look for 'Data' or 'Capture' in names
            # 2. Or check for valid headers in first few sheets
            target_sheet = xls.sheet_names[0]
            best_sheet = None
            max_score = 0
            
            for sheet in xls.sheet_names:
                try:
                     # Peek at first few rows
                     df_check = pd.read_excel(xlsx, sheet, nrows=5)
                     cols_str = " ".join([str(c).lower() for c in df_check.columns])
                     
                     score = 0
                     if 'timestamp' in cols_str or 'date' in cols_str: score += 5
                     if 'watt' in cols_str or 'kw' in cols_str or 'power' in cols_str: score += 5
                     if 'captured' in sheet.lower() or 'data' in sheet.lower(): score += 2
                     
                     if score > max_score and score >= 5:
                         max_score = score
                         best_sheet = sheet
                except:
                     continue
            
            if best_sheet:
                 target_sheet = best_sheet
                 st.info(f"Selecting sheet '{target_sheet}' based on content.")
            
            df = pd.read_excel(xlsx, target_sheet)
            
            # Check for header issue: if columns are "Unnamed" and first row looks like headers
            # A simple heuristic: check if 'localtime' or 'watt' or 'date' is in the first row values
            first_row_vals = [str(x).lower() for x in df.iloc[0].values] if not df.empty else []
            if any('localtime' in x or 'watt' in x or 'date' in x for x in first_row_vals):
                # Reload with header=1
                file.seek(0)
                df = pd.read_excel(xlsx, sheet_name=target_sheet, header=1)
                st.info("Detected header in second row. Reloaded file.")

        else:
            # Assume CSV
            # Check for BOM
            file.seek(0)
            bom = file.read(2)
            file.seek(0)
            
            detected_encoding = None
            if bom == b'\xff\xfe':
                detected_encoding = 'utf-16-le'
            elif bom == b'\xfe\xff':
                detected_encoding = 'utf-16-be'
            
            encodings = [detected_encoding] if detected_encoding else ['utf-8', 'utf-16', 'utf-16-le', 'ISO-8859-1', 'cp1252']
            
            for enc in encodings:
                try:
                    file.seek(0)
                    df = pd.read_csv(file, encoding=enc)
                    
                    # Heuristic for valid load: check for expected keywords
                    cols_clean = [str(c).strip().lstrip('\ufeff').lower() for c in df.columns]
                    
                    keywords = ['date', 'time', 'timestamp', 'localtime', 'watt', 'kw', 'power']
                    if any(k in cols_clean for k in keywords):
                         break
                    
                    # Also keep it if it looks visibly cleaner than previous?
                    # For now, keyword match is strong signal.
                except Exception:
                    continue
            else:
                 # If we exhausted specific encodings, warn, but maybe we got a DF from last attempt?
                 # No, loop var scope.
                 pass
            
            if 'df' not in locals() or df is None:
                 st.error("Failed to read CSV with common encodings.")
                 return None, None

    except Exception as e:
        st.error(f"Error reading file: {e}")
        return None, None

    # Column Identification
    col_map = {}
    
    # Normalize columns: strip whitespace and lower case, ensure string
    # Remove BOM artifacts if any remain in column names
    df.columns = [str(c).strip().lstrip('\ufeff').lower() for c in df.columns]
    columns_lower = list(df.columns)

    # 0. Check for separate Date and Time columns
    if 'date' in columns_lower and 'time' in columns_lower:
        try:
             # Combine them
             date_col = df.columns[columns_lower.index('date')]
             time_col = df.columns[columns_lower.index('time')]
             df['Combined_Timestamp'] = df[date_col].astype(str) + ' ' + df[time_col].astype(str)
             col_map['timestamp'] = 'Combined_Timestamp'
             st.info(f"Merged '{date_col}' and '{time_col}' into Timestamp.")
        except Exception as e:
             st.warning(f"Found Date and Time but failed to merge: {e}")

    # Timestamp identification priority:
    # 1. 'localtime' (explicit user mention)
    # 2. 'Date/Time' (common standard)
    # 3. 'timestamp'
    # 4. Any column with 'date' or 'time'
    
    if 'timestamp' not in col_map:
        if 'localtime' in columns_lower:
            col_map['timestamp'] = df.columns[columns_lower.index('localtime')]
        elif 'date/time' in columns_lower:
            col_map['timestamp'] = df.columns[columns_lower.index('date/time')]
        elif 'timestamp' in columns_lower:
            col_map['timestamp'] = df.columns[columns_lower.index('timestamp')]
        else:
            # Fallback to substring search
            time_cols = [c for c in df.columns if 'time' in str(c).lower() or 'date' in str(c).lower()]
            if time_cols:
                col_map['timestamp'] = time_cols[0]
            else:
                st.error("Could not identify a Timestamp column (looking for 'Date' + 'Time', 'localtime', 'Date/Time', 'timestamp', etc.).")
                return None, None
            
    st.info(f"Identified Timestamp Column: {col_map['timestamp']}")

    # Power
    # KW Mapping Priority for NSIA and others:
    # 1. 'active power total avg' (NSIA - Aggregated Avg)
    # 2. 'active power total'
    # 3. 'kw'
    # 4. 'total active power'
    # 5. 'active power'
    # 6. 'watt' 
    potential_kw = ['active power total avg', 'active power total', 'active power total max', 
                    'kw', 'total active power', 'active power', 'watt', 'power']
    
    found_power = None
    for p in potential_kw:
        # Check against lowercased columns in loop order
        # We need exact match or close substring match?
        # Use simple substring check, but prioritized by order in list
        match = next((c for c in columns_lower if p in c), None)
        if match:
             # Sanity check to avoid partial matches on wrong things (e.g. 'reactive power' matching 'power')
             # 'active power' in list prevents 'reactive' issue if logic handles it?
             # 'power' is last resort.
             if 'reactive' in match and 'reactive' not in p:
                  continue # Skip if we matched 'reactive power' but wanted 'power' or 'active power'
             if 'apparent' in match and 'apparent' not in p:
                  continue
             
             found_power = match
             break
    
    if found_power:
        col_map['power'] = found_power # We use 'power' key temporarily before renaming to 'kw'
        col_map['kw'] = found_power # Set 'kw' key directly to be safe
        st.info(f"Identified Power Column: {col_map['kw']}")
    else:
        st.error("Could not identify a Power column (looking for 'watt', 'active power', 'kw').")
        return None, None

    # Power Factor
    pf_cols = [c for c in df.columns if 'pf' in str(c).lower() or 'factor' in str(c).lower()]
    if pf_cols:
        col_map['pf'] = pf_cols[0]
    else:
        st.warning("Could not identify a Power Factor column. Assuming PF=1.0 for now.")
        col_map['pf'] = None


    # Cleaning and Conversion
    try:
        # Attempt standard parsing, prioritizing dayfirst for DD.MM.YYYY formats common in CSVs
        df[col_map['timestamp']] = pd.to_datetime(df[col_map['timestamp']], dayfirst=True, errors='coerce')
        
        # If too many NaTs (e.g., > 90% failure), maybe it's a unix timestamp?
        # But 'localtime' implies string. 'timestamp' implies unix.
        # If we picked 'timestamp' column which is numeric:
        if df[col_map['timestamp']].isna().mean() > 0.9:
            # Check if it looks like unix epoch (numeric)
            # Re-read or just check the raw values? 
            # It's hard to rollback the .dt conversion in-place easily without reloading or keeping a copy.
            # But wait, to_datetime with errors='coerce' turns non-parsable to NaT.
            # If we selected 'timestamp' (unix), to_datetime might have failed if it expected strings or if numbers were large? 
            # Actually pd.to_datetime on unix int usually works if unit is specified, otherwise it assumes nanoseconds (default) 
            # which for 1714113225 (seconds) -> 1970ish. 
            # Let's not over-engineer unless needed. The user emphasized 'localtime'.
            pass
            
    except Exception as e:
        st.error(f"Error parsing timestamp: {e}")
        return None, None

    # Convert to numeric, forcing errors to NaN
    df[col_map['power']] = pd.to_numeric(df[col_map['power']], errors='coerce')
    if col_map['pf']:
        df[col_map['pf']] = pd.to_numeric(df[col_map['pf']], errors='coerce')
    
    df = df.dropna(subset=[col_map['power'], col_map['timestamp']])

    # Unit Conversion logic
    max_power = df[col_map['power']].max()
    if max_power > 2000:
        df['kW'] = df[col_map['power']] / 1000.0
        st.info(f"Detected Power in Watts (Max: {max_power:.2f}). Converted to kW.")
    else:
        df['kW'] = df[col_map['power']]
        st.info(f"Detected Power in kW (Max: {max_power:.2f}). No conversion needed.")
    
    # Rename for consistency
    rename_dict = {col_map['timestamp']: 'Timestamp', col_map['pf']: 'Power Factor'} if col_map['pf'] else {col_map['timestamp']: 'Timestamp'}
    df = df.rename(columns=rename_dict)
    
    return df, col_map

def calculate_metrics(df):
    """
    Aggregates data and calculates required metrics using interval-based logic.
    """
    df['Timestamp'] = pd.to_datetime(df['Timestamp'])
    df['Day'] = df['Timestamp'].dt.date # Use full date to handle multi-month data correctly
    df['Hour'] = df['Timestamp'].dt.hour
    
    cols_actual = {c.lower(): c for c in df.columns}
    
    # Generic Phase Detection
    # Config: List of (PhaseA, PhaseB, PhaseC) candidates
    # We check each triplet. If all 3 exist, we use them.
    phase_candidates = [
        ('watt_a', 'watt_b', 'watt_c'), # Pirano, Transformer 1 (likely)
        ('active power l1n avg', 'active power l2n avg', 'active power l3n avg'), # NSIA
        ('active power l1', 'active power l2', 'active power l3'), # Generic
        ('power l1', 'power l2', 'power l3')
    ]
    
    found_triplet = None
    for p_a, p_b, p_c in phase_candidates:
         # Check strict existence in lowercased columns
         if p_a in cols_actual and p_b in cols_actual and p_c in cols_actual:
              found_triplet = (p_a, p_b, p_c)
              break
    
    # If not found strict, try partial match? (Risky)
    # Let's stick to known schemas.
    
    if found_triplet:
         # Only override if we really want to?
         # User said "check this 3 phase system... aggregate to get maximum power"
         # For Pirano, we did it.
         # For NSIA, we might have 'active power total avg' mapped to kw.
         # If we aggregate L1+L2+L3, we expect same result.
         # But if the file is "filled with errors", recalculating might be safer?
         # "Most files are like that".
         # Let's Apply Aggregation if found, potentially overwriting 'kw' (or 'Total_Watt_Phases').
         
         # Wait, if we overwrite 'kw' which was 'active power total avg', we are asserting Sum(Phases) is truth.
         # This is generally safe for Active Power.
         
         p_a, p_b, p_c = found_triplet
         w_a = df[cols_actual[p_a]].fillna(0)
         w_b = df[cols_actual[p_b]].fillna(0)
         w_c = df[cols_actual[p_c]].fillna(0)
         
         # Unit detection:
         # If column has 'watt' in name -> Watts. If 'kw' -> kW.
         # NSIA: 'active power l1n avg' (usually Watts/kW?). Value check needed?
         # Pirano: 'watt_a' -> Watts.
         # We'll assume Watts if 'watt' in name, else check values?
         # For NSIA, 'active power l1n avg' is likely kW? Or W?
         # If we divide by 1000 blindly, we might underreport.
         # Heuristic: If max value > 10000, assumes Watts. If < 5000, assumes kW (unless big factory).
         # Pirano was ~10000 -> 30kW.
         # Let's try to infer from name OR value.
         
         total_raw = w_a + w_b + w_c
         
         if 'watt' in p_a or 'watt' in p_b:
             # Assume Watts
             df['kW'] = total_raw / 1000.0
             st.info(f"Aggregated 3-Phase Power ({p_a}+{p_b}+{p_c}) [Watts -> kW].")
         else:
             # Assume kW (e.g. 'active power' usually kw)
             df['kW'] = total_raw
             st.info(f"Aggregated 3-Phase Power ({p_a}+{p_b}+{p_c}) [Summed].")
         
    # Detect Interval
    # ...
    # Calculate median time difference in minutes
    if len(df) > 1:
        # distinct timestamps
        unique_times = df['Timestamp'].sort_values().unique()
        if len(unique_times) > 1:
            diffs = pd.Series(unique_times).diff().dt.total_seconds() / 60.0
            median_interval_min = diffs.median()
        else:
            median_interval_min = 60 # Default to hourly if only 1 timestamp? or 0?
    else:
        median_interval_min = 60 # Default

    st.info(f"Detected Data Interval: {median_interval_min:.2f} minutes")

    if median_interval_min < 55: # Tolerance for 1-hour
        # Sub-hourly: 
        # User Request: "moving ahead to get average per hour"
        # Previous logic: 'Accumulated Power' (Sum). New logic: 'Average Power' (Mean).
        agg_method = 'mean'
        metric_label = 'Average Power (kW)'
        
        # Mean of kW for each specific hour
        full_hourly_profile = df.groupby(['Day', 'Hour'])['kW'].mean()
        
        # Baseload: Average of the hourly MEANS
        baseload_subset = full_hourly_profile.reset_index()
        baseload_kw = baseload_subset[baseload_subset['Hour'].isin([1,2,3,4])]['kW'].mean()
        
    else:
        # Hourly (or greater): Use MEAN (Average Inst. Power)
        # If it's truly hourly, Mean = The Value itself.
        agg_method = 'mean'
        metric_label = 'Average Power (kW)'
        full_hourly_profile = df.groupby(['Day', 'Hour'])['kW'].mean()
        
        # Baseload: Average of the hourly MEANS
        baseload_subset = full_hourly_profile.reset_index()
        baseload_kw = baseload_subset[baseload_subset['Hour'].isin([1,2,3,4])]['kW'].mean()

    # Calculate Global Peak
    global_peak_kw = full_hourly_profile.max()
    
    # Find timestamp of peak
    peak_day, peak_hour = full_hourly_profile.idxmax()
    
    # Approximate timestamp (first entry matching Day/Hour)
    peak_timestamp_rows = df[(df['Day'] == peak_day) & (df['Hour'] == peak_hour)]
    peak_timestamp = peak_timestamp_rows['Timestamp'].iloc[0] if not peak_timestamp_rows.empty else None

    avg_pf = df['Power Factor'].mean() if 'Power Factor' in df.columns else None

    # For the chart, we keep the 0-23 Average Profile shape
    avg_hourly_profile = full_hourly_profile.groupby('Hour').mean()

    return {
        'hourly_profile': avg_hourly_profile, 
        'full_profile': full_hourly_profile,
        'global_peak_kw': global_peak_kw,
        'peak_timestamp': peak_timestamp,
        'avg_pf': avg_pf,
        'baseload_kw': baseload_kw,
        'agg_method': agg_method,
        'metric_label': metric_label
    }

def generate_chart(hourly_profile, label="Hourly Load Profile"):
    """
    Generates Matplotlib figure for the hourly profile.
    """
    plt.style.use('ggplot')
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.plot(hourly_profile.index, hourly_profile.values, marker='o', linestyle='-')
    ax.set_title(label)
    ax.set_xlabel("Hour of Day")
    ax.set_ylabel("Power (kW)")
    ax.set_xticks(range(0, 24))
    ax.grid(True)
    return fig

def generate_excel(df_raw, hourly_profile, agg_method='sum'):
    """
    Generates an Excel file with raw data, formulas, and charts.
    """
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        # Write Raw Data
        df_raw.to_excel(writer, sheet_name='Raw_Analysis', index=False)
        workbook = writer.book
        worksheet_raw = writer.sheets['Raw_Analysis']
        worksheet_profile = workbook.add_worksheet('Hourly_Profiles')
        
        cols = df_raw.columns.tolist()
        
        def get_col_letter(name):
            try:
                idx = cols.index(name)
                return chr(65 + idx)
            except ValueError:
                return None

        col_day = get_col_letter('Day')
        col_hour = get_col_letter('Hour')
        col_kw = get_col_letter('kW')
        
        # Determine Label and Formula based on Agg Method
        if agg_method == 'sum':
            prof_header = 'Accumulated Power (kW)'
            formula_base = 'SUMIFS'
        else:
            prof_header = 'Average Power (kW)'
            formula_base = 'AVERAGEIFS'

        # Write Headers for Profile
        worksheet_profile.write('A1', 'Day')
        worksheet_profile.write('B1', 'Hour')
        worksheet_profile.write('C1', prof_header)
        
        if col_day and col_hour and col_kw:
            desc_format = workbook.add_format({'num_format': '0.00'})
            date_format = workbook.add_format({'num_format': 'yyyy-mm-dd'})
            unique_days = sorted(df_raw['Day'].unique())
            row = 2
            for day in unique_days:
                for h in range(24):
                    worksheet_profile.write(f'A{row}', day, date_format)
                    worksheet_profile.write(f'B{row}', h)
                    
                    if agg_method == 'sum':
                         # SUMIFS(kW, Day_col, day, Hour_col, h)
                         formula = f'=SUMIFS(Raw_Analysis!{col_kw}:{col_kw}, Raw_Analysis!{col_day}:{col_day}, A{row}, Raw_Analysis!{col_hour}:{col_hour}, B{row})'
                    else:
                         # AVERAGEIFS(kW, Day_col, day, Hour_col, h)
                         formula = f'=AVERAGEIFS(Raw_Analysis!{col_kw}:{col_kw}, Raw_Analysis!{col_day}:{col_day}, A{row}, Raw_Analysis!{col_hour}:{col_hour}, B{row})'

                    worksheet_profile.write_formula(f'C{row}', formula, desc_format)
                    row += 1
        else:
            worksheet_profile.write('C1', 'Error: Could not locate data columns for formulas.')
            hourly_profile.to_excel(writer, sheet_name='Hourly_Profiles_Backup')

        # --- Add Chart ---
        chart = workbook.add_chart({'type': 'line'})
        
        worksheet_profile.write('H1', 'Hour')
        worksheet_profile.write('I1', 'Avg ' + prof_header)
        
        for i in range(24):
            r = i + 2
            worksheet_profile.write(f'H{r}', i)
            worksheet_profile.write_formula(f'I{r}', f'=AVERAGEIF(B:B, H{r}, C:C)', desc_format)

        chart.add_series({
            'name':       'Average ' + prof_header,
            'categories': f'=Hourly_Profiles!$H$2:$H$25',
            'values':     f'=Hourly_Profiles!$I$2:$I$25',
            'line':       {'color': 'blue'},
        })
        
        chart.set_title ({'name': 'Daily Load Profile'})
        chart.set_x_axis({'name': 'Hour of Day'})
        chart.set_y_axis({'name': 'Power (kW)'})
        chart.set_style(10)
        
        worksheet_profile.insert_chart('D2', chart)

    return output.getvalue()

def get_ai_analysis(metrics, api_key):
    """
    Generates technical observations using Gemini.
    """
    if not api_key:
        return "API Key missing."
    
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel('gemini-2.5-flash')
    
    prompt = f"""
    You are a Technical Data Analyst.
    Write a direct, concise technical analysis (approx. 400-500 words) based on the metrics below.
    
    IMPORTANT CONTEXT: 
    The "Global Peak" and load profile values represent the ACCUMULATED hourly load (Sum of minute-interval readings), NOT the average instantaneous power. This explains why the values (e.g., >500 kW) may appear significantly higher than the facility's rated capacity (e.g., 50 kVA). Do not flag this as an error; interpret it as total hourly energy intensity.

    Constraints:
    - NO em dashes (—). Use colons or parentheses if needed.
    - NO conversational filler (e.g., "It is worth noting that...").
    - Start sentences directly with the subject.
    - Use active voice.
    
    Structure:
    1. Load Profile: Consumption pattern, peak timing, baseload intensity.
    2. Power Factor: Efficiency implications.
    3. Anomalies: Spikes or dips (relative to the profile).
    4. Implications: System health and energy intensity.
    
    Metrics:
    - Global Hourly Accumulated Peak: {metrics['global_peak_kw']:.2f} (Max of Hourly Sums) at {metrics['peak_timestamp']}
    - Average Power Factor: {metrics['avg_pf']:.2f}
    - Nighttime Accumulated Baseload (01:00-04:00): {metrics['baseload_kw']:.2f} (Avg of Hourly Sums)
    """
    
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error generating AI analysis: {e}"

def generate_word_report(metrics, chart_fig, ai_text):
    """
    Generates a Word document.
    """
    doc = Document()
    doc.add_heading('Energy Audit Report', 0)
    
    # Removed Letter headers (Date, To, From) per user request
    
    doc.add_heading('System Metrics', level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Global Accumulated Peak'
    row_cells[1].text = f"{metrics['global_peak_kw']:.2f} kW"
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Peak Timestamp'
    row_cells[1].text = str(metrics['peak_timestamp'])
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Average Power Factor'
    row_cells[1].text = f"{metrics['avg_pf']:.2f}"
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Nighttime Accumulated Baseload'
    row_cells[1].text = f"{metrics['baseload_kw']:.2f} kW"
    
    doc.add_heading('Daily Load Profile', level=1)
    
    # Save plot to buffer
    img_buffer = io.BytesIO()
    chart_fig.savefig(img_buffer, format='png', dpi=300)
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(6.0))
    
    doc.add_heading('Technical Analysis', level=1)
    doc.add_paragraph(ai_text)
    
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()
